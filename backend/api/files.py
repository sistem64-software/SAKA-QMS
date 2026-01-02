from fastapi import APIRouter, HTTPException, UploadFile, Form
from fastapi.responses import FileResponse
from pathlib import Path
from typing import List, Dict, Any
import openpyxl
from docx import Document
import json
import shutil
import os
import zipfile
import xml.etree.ElementTree as ET

router = APIRouter()

TEMPLATE_DIR = Path("form_sablonlari")
COMPANIES_DIR = Path("firmalar")
COMPANIES_DIR.mkdir(exist_ok=True)

@router.get("/file/{filename}")
async def get_file(filename: str):
    """
    Şablon dosyasını oku (parse edilmiş JSON formatında)
    """
    file_path = TEMPLATE_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Dosya bulunamadı")
    
    try:
        extension = file_path.suffix.lower()
        
        if extension in ['.xlsx', '.xls']:
            return read_excel(file_path)
        elif extension in ['.docx', '.doc']:
            return read_word(file_path)
        else:
            raise HTTPException(status_code=400, detail="Desteklenmeyen dosya formatı")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya okuma hatası: {str(e)}")

@router.get("/file/{filename}/raw")
async def get_file_raw(filename: str):
    """
    Şablon dosyasını ham haliyle (raw) döndür - Syncfusion için
    """
    file_path = TEMPLATE_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Dosya bulunamadı")
    
    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

@router.get("/companies/{company_name}/file/{filename}")
async def get_company_file(company_name: str, filename: str):
    """
    Firma dosyasını oku (parse edilmiş JSON formatında, alt klasörler dahil)
    """
    company_dir = COMPANIES_DIR / company_name
    
    if not company_dir.exists():
        raise HTTPException(status_code=404, detail="Firma bulunamadı")
    
    # Dosyayı firma klasöründe veya alt klasörlerde ara
    file_path = None
    for found_path in company_dir.rglob(filename):
        if found_path.is_file() and found_path.name == filename:
            file_path = found_path
            break
    
    if not file_path or not file_path.exists():
        raise HTTPException(status_code=404, detail="Dosya bulunamadı")
    
    try:
        extension = file_path.suffix.lower()
        
        if extension in ['.xlsx', '.xls']:
            return read_excel(file_path)
        elif extension in ['.docx', '.doc']:
            return read_word(file_path)
        else:
            raise HTTPException(status_code=400, detail="Desteklenmeyen dosya formatı")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya okuma hatası: {str(e)}")

@router.get("/companies/{company_name}/file/{filename}/raw")
async def get_company_file_raw(company_name: str, filename: str):
    """
    Firma dosyasını ham haliyle (raw) döndür - Syncfusion için (alt klasörler dahil)
    """
    company_dir = COMPANIES_DIR / company_name
    
    if not company_dir.exists():
        raise HTTPException(status_code=404, detail="Firma bulunamadı")
    
    # Dosyayı firma klasöründe veya alt klasörlerde ara
    file_path = None
    for found_path in company_dir.rglob(filename):
        if found_path.is_file() and found_path.name == filename:
            file_path = found_path
            break
    
    if not file_path or not file_path.exists():
        raise HTTPException(status_code=404, detail="Dosya bulunamadı")
    
    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

@router.delete("/file/{filename}")
async def delete_file(filename: str):
    """
    Şablon dosyasını sil
    """
    file_path = TEMPLATE_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Dosya bulunamadı")
    
    try:
        os.remove(file_path)
        return {
            "message": "Dosya başarıyla silindi",
            "filename": filename
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya silme hatası: {str(e)}")

@router.delete("/companies/{company_name}/file/{filename}")
async def delete_company_file(company_name: str, filename: str):
    """
    Firma dosyasını sil (alt klasörler dahil)
    """
    company_dir = COMPANIES_DIR / company_name
    
    if not company_dir.exists():
        raise HTTPException(status_code=404, detail="Firma bulunamadı")
    
    # Dosyayı firma klasöründe veya alt klasörlerde ara
    file_path = None
    for found_path in company_dir.rglob(filename):
        if found_path.is_file() and found_path.name == filename:
            file_path = found_path
            break
    
    if not file_path or not file_path.exists():
        raise HTTPException(status_code=404, detail="Dosya bulunamadı")
    
    try:
        os.remove(file_path)
        return {
            "message": "Dosya başarıyla silindi",
            "company": company_name,
            "filename": filename
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya silme hatası: {str(e)}")

@router.patch("/file/rename")
async def rename_file(old_name: str, new_name: str):
    """
    Şablon dosyasının adını değiştir
    """
    old_path = TEMPLATE_DIR / old_name
    
    if not old_path.exists():
        raise HTTPException(status_code=404, detail="Dosya bulunamadı")
    
    # Yeni dosya adı için uzantıyı koru veya kontrol et
    old_ext = old_path.suffix.lower()
    if not new_name.lower().endswith(old_ext):
        new_name += old_ext
    
    new_path = TEMPLATE_DIR / new_name
    
    if new_path.exists():
        raise HTTPException(status_code=400, detail="Bu isimde bir dosya zaten mevcut")
    
    try:
        os.rename(old_path, new_path)
        return {
            "message": "Dosya başarıyla adlandırıldı",
            "old_name": old_name,
            "new_name": new_name
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya adlandırma hatası: {str(e)}")

@router.patch("/companies/{company_name}/file/rename")
async def rename_company_file(company_name: str, old_name: str, new_name: str):
    """
    Firma dosyasının adını değiştir (alt klasörler dahil)
    """
    company_dir = COMPANIES_DIR / company_name
    
    if not company_dir.exists():
        raise HTTPException(status_code=404, detail="Firma bulunamadı")
    
    # Dosyayı firma klasöründe veya alt klasörlerde ara
    old_path = None
    for found_path in company_dir.rglob(old_name):
        if found_path.is_file() and found_path.name == old_name:
            old_path = found_path
            break
    
    if not old_path or not old_path.exists():
        raise HTTPException(status_code=404, detail="Dosya bulunamadı")
    
    # Yeni dosya adı için uzantıyı koru veya kontrol et
    old_ext = old_path.suffix.lower()
    if not new_name.lower().endswith(old_ext):
        new_name += old_ext
    
    new_path = old_path.parent / new_name
    
    if new_path.exists():
        raise HTTPException(status_code=400, detail="Bu isimde bir dosya zaten mevcut")
    
    try:
        os.rename(old_path, new_path)
        return {
            "message": "Dosya başarıyla adlandırıldı",
            "company": company_name,
            "old_name": old_name,
            "new_name": new_name
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya adlandırma hatası: {str(e)}")


@router.post("/save-word-file")
async def save_word_file_upload(file: UploadFile, company: str = Form(...)):
    """
    Düzenlenmiş Word dosyasını firma klasörüne kaydet
    Syncfusion'dan gelen .docx blob'unu alır ve firma klasörüne kaydeder
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Dosya adı geçersiz")
    
    if not company:
        raise HTTPException(status_code=400, detail="Firma adı belirtilmedi")
    
    # Firma klasörünü oluştur
    company_dir = COMPANIES_DIR / company
    company_dir.mkdir(parents=True, exist_ok=True)
    
    # Dosya adına firma prefix'i ekle
    company_prefix = company.upper().replace(' ', '_')
    new_filename = f"{company_prefix}_{file.filename}"
    
    # Hedef dosya yolu
    target_path = company_dir / new_filename
    
    try:
        # Dosyayı kaydet
        with open(target_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        return {
            "message": "Dosya başarıyla kaydedildi",
            "path": str(target_path),
            "company": company,
            "filename": new_filename,
            "size": os.path.getsize(target_path)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya kaydetme hatası: {str(e)}")
    finally:
        file.file.close()


def read_excel(file_path: Path) -> Dict[str, Any]:
    """Excel dosyasını oku ve tüm format bilgileriyle JSON formatına çevir"""
    import base64
    from io import BytesIO
    
    wb = openpyxl.load_workbook(file_path)
    
    sheets_data = {}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        
        # Sheet'teki tüm verileri ve stilleri al
        data = []
        for row in ws.iter_rows():
            row_data = []
            for cell in row:
                # Hücre değeri
                value = cell.value if cell.value is not None else ""
                
                # Hücre stilleri
                cell_style = {
                    "value": str(value),
                    "coordinate": cell.coordinate,
                }
                
                # Formül varsa ekle (data_type 'f' ise formül hücresidir)
                if cell.data_type == 'f':
                    cell_style["formula"] = cell.value if isinstance(cell.value, str) and cell.value.startswith('=') else f"={cell.value}"
                    # Hesaplanmış değeri de sakla (eğer varsa)
                    try:
                        # Formülün hesaplanmış değerini al
                        calculated_value = ws.cell(cell.row, cell.column).value
                        if calculated_value is not None:
                            cell_style["value"] = str(calculated_value)
                    except:
                        pass
                
                # Font stilleri
                if cell.font:
                    cell_style["font"] = {
                        "bold": cell.font.bold or False,
                        "italic": cell.font.italic or False,
                        "size": cell.font.size or 11,
                        "color": cell.font.color.rgb if cell.font.color and hasattr(cell.font.color, 'rgb') else None,
                        "name": cell.font.name or "Calibri"
                    }
                
                # Arka plan rengi
                if cell.fill and cell.fill.start_color:
                    if hasattr(cell.fill.start_color, 'rgb'):
                        cell_style["fill"] = cell.fill.start_color.rgb
                
                # Hizalama
                if cell.alignment:
                    cell_style["alignment"] = {
                        "horizontal": cell.alignment.horizontal or "left",
                        "vertical": cell.alignment.vertical or "top",
                        "wrap_text": cell.alignment.wrap_text or False
                    }
                
                # Kenarlıklar
                if cell.border:
                    cell_style["border"] = {
                        "top": bool(cell.border.top and cell.border.top.style),
                        "bottom": bool(cell.border.bottom and cell.border.bottom.style),
                        "left": bool(cell.border.left and cell.border.left.style),
                        "right": bool(cell.border.right and cell.border.right.style)
                    }
                
                row_data.append(cell_style)
            data.append(row_data)
        
        # Birleştirilmiş hücreler
        merged_cells = []
        for merged_range in ws.merged_cells.ranges:
            merged_cells.append({
                "start_row": merged_range.min_row,
                "start_col": merged_range.min_col,
                "end_row": merged_range.max_row,
                "end_col": merged_range.max_col,
                "range": str(merged_range)
            })
        
        # Sütun genişlikleri
        column_widths = {}
        for col_letter, col_dimension in ws.column_dimensions.items():
            if col_dimension.width:
                column_widths[col_letter] = col_dimension.width
        
        # Satır yükseklikleri
        row_heights = {}
        for row_num, row_dimension in ws.row_dimensions.items():
            if row_dimension.height:
                row_heights[str(row_num)] = row_dimension.height
        
        # Resimleri oku
        images = []
        
        # Daha önce eklenen resimlerin anchor'larını takip et (Deep Parse ile çakışmayı önlemek için)
        existing_anchors = set()
        
        # 1. Önce Deep Parse (Hücre İçi) resimleri dene
        try:
            deep_images = extract_deep_images(file_path)
            if sheet_name in deep_images:
                print(f"Deep parse ile {sheet_name} sayfasında {len(deep_images[sheet_name])} resim bulundu.")
                for img in deep_images[sheet_name]:
                    images.append(img)
                    if img.get('anchor'):
                        existing_anchors.add(img['anchor'])
        except Exception as e:
            print(f"Deep parse hatası: {str(e)}")

        # 2. Standart OpenPyXL Resimleri
        if hasattr(ws, '_images') and ws._images:
            for img in ws._images:
                try:
                    # Resim verisini al
                    image_data = img._data()
                    
                    # Base64'e çevir
                    base64_data = base64.b64encode(image_data).decode('utf-8')
                    
                    # Resim formatını belirle
                    image_format = 'png'  # Varsayılan
                    if hasattr(img, 'format'):
                        image_format = img.format.lower()
                    
                    # Anchor bilgisini al (resmin bağlı olduğu hücre)
                    # Anchor bilgisini al (resmin bağlı olduğu hücre)
                    anchor = None
                    try:
                        if hasattr(img, 'anchor'):
                            # Farklı anchor tiplerini kontrol et
                            # 1. TwoCellAnchor veya OneCellAnchor (genelde _from özelliğine sahiptir)
                            if hasattr(img.anchor, '_from'):
                                anchor_from = img.anchor._from
                                col = getattr(anchor_from, 'col', None)
                                row = getattr(anchor_from, 'row', None)
                                
                                if col is not None and row is not None:
                                    col_letter = openpyxl.utils.get_column_letter(col + 1)
                                    row_num = row + 1
                                    anchor = f"{col_letter}{row_num}"
                            
                            # 2. Eğer yukarıdaki yöntem çalışmadıysa ve direkt col/row varsa (bazen AbsoluteAnchor olabilir ama hücreye bağlı değildir)
                            elif not anchor and hasattr(img.anchor, 'col') and hasattr(img.anchor, 'row'):
                                col = img.anchor.col
                                row = img.anchor.row
                                col_letter = openpyxl.utils.get_column_letter(col + 1)
                                row_num = row + 1
                                anchor = f"{col_letter}{row_num}"
                                
                            # 3. String olarak gelme ihtimali (eski sürümler vs)
                            elif isinstance(img.anchor, str):
                                anchor = img.anchor
                                
                        # Eğer bu anchor zaten Deep Parse ile bulunduysa, tekrar ekleme
                        # (Genelde Place in Cell resimleri openpyxl'de görünmez ama yine de çakışma kontrolü)
                        if anchor and anchor in existing_anchors:
                            continue

                    except Exception as anchor_err:
                        print(f"Anchor okuma hatası: {str(anchor_err)}")
                        # Anchor okunamadıysa resmi formatsız ekle veya atla
                        pass
                    
                    # Resim boyutları
                    width = img.width if hasattr(img, 'width') else None
                    height = img.height if hasattr(img, 'height') else None
                    
                    images.append({
                        "anchor": anchor,
                        "data": base64_data,
                        "format": image_format,
                        "width": width,
                        "height": height
                    })
                except Exception as e:
                    print(f"Resim okuma hatası: {str(e)}")
                    continue
        
        sheets_data[sheet_name] = {
            "data": data,
            "max_row": ws.max_row,
            "max_column": ws.max_column,
            "merged_cells": merged_cells,
            "column_widths": column_widths,
            "row_heights": row_heights,
            "images": images
        }
    
    return {
        "type": "excel",
        "filename": file_path.name,
        "sheets": sheets_data,
        "active_sheet": wb.active.title
    }

def read_word(file_path: Path) -> Dict[str, Any]:
    """Word dosyasını oku ve JSON formatına çevir - stil bilgileriyle birlikte"""
    doc = Document(file_path)
    
    # Paragrafları oku
    paragraphs = []
    for para in doc.paragraphs:
        para_data = {
            "text": para.text,
            "style": para.style.name if para.style else "Normal"
        }
        
        # Paragraph format bilgileri
        if para.runs:
            # İlk run'un formatını al (genelde tüm paragraph aynı formatta)
            first_run = para.runs[0]
            para_data["format"] = {
                "bold": first_run.bold if first_run.bold is not None else False,
                "italic": first_run.italic if first_run.italic is not None else False,
                "underline": bool(first_run.underline),
                "font_size": first_run.font.size.pt if first_run.font.size else None,
                "font_name": first_run.font.name if first_run.font.name else None,
                "font_color": str(first_run.font.color.rgb) if first_run.font.color and hasattr(first_run.font.color, 'rgb') else None
            }
            
            # Alignment
            if para.alignment:
                para_data["alignment"] = str(para.alignment)
        
        paragraphs.append(para_data)
    
    # Tabloları oku
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_data = {
                    "text": cell.text
                }
                # Cell format bilgileri
                if cell.paragraphs:
                    first_para = cell.paragraphs[0]
                    if first_para.runs:
                        first_run = first_para.runs[0]
                        cell_data["format"] = {
                            "bold": first_run.bold if first_run.bold is not None else False,
                            "italic": first_run.italic if first_run.italic is not None else False,
                            "font_size": first_run.font.size.pt if first_run.font.size else None,
                            "font_name": first_run.font.name if first_run.font.name else None,
                            "font_color": str(first_run.font.color.rgb) if first_run.font.color and hasattr(first_run.font.color, 'rgb') else None
                        }
                row_data.append(cell_data)
            table_data.append(row_data)
        tables.append(table_data)
    
    return {
        "type": "word",
        "filename": file_path.name,
        "paragraphs": paragraphs,
        "tables": tables
    }

def sanitize_folder_name(name: str) -> str:
    """
    Klasör adını temizle - geçersiz karakterleri kaldır
    """
    import re
    name = re.sub(r'[<>:"/\\|?*]', '_', str(name))
    name = name.strip()
    return name if name else None

def extract_work_order_number(sheets_data: Dict[str, Any]) -> str:
    """
    Excel içeriğinden iş emri numarasını çıkar
    SM ile başlayan herhangi bir hücreyi iş emri no olarak kabul eder
    Fixed: Removed emoji characters for Windows compatibility
    """
    import re
    
    print(f"\n====== EXTRACT_WORK_ORDER_NUMBER (SM ile başlayan hücre aranıyor) ======")
    print(f"Sheets: {list(sheets_data.keys())}")
    
    try:
        for sheet_name, sheet in sheets_data.items():
            data = sheet.get('data', [])
            
            print(f"\n--- Sheet: {sheet_name} ---")
            
            for row_idx, row in enumerate(data):
                for col_idx, cell in enumerate(row):
                    cell_value = str(cell.get('value', '')).strip()
                    
                    if not cell_value:
                        continue
                    
                    # SM ile başlıyor mu?
                    if cell_value.upper().startswith('SM'):
                        print(f"[OK] Is Emri NO bulundu: Row {row_idx}, Col {col_idx} = '{cell_value}'")
                        sanitized = sanitize_folder_name(cell_value)
                        if sanitized:
                            print(f"   [INFO] IS EMRI NO: '{sanitized}'")
                            return sanitized
        
        print("[WARN] SM ile baslayan hucre bulunamadi")
        return None
    except Exception as e:
        print(f"[ERROR] Hata: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

@router.post("/save")
async def save_file(data: Dict[str, Any]):
    """
    Düzenlenmiş dosyayı firma klasörüne kaydet
    Dosya adı başına firma ismi eklenir (örn: BAYKAR_F.02.xlsx)
    İş emri no bulunursa alt klasör oluşturulur (örn: firmalar/Baykar/SM-128/)
    """
    try:
        filename = data.get("filename")
        company_name = data.get("company")
        file_type = data.get("type")
        content = data.get("content")
        
        if not all([filename, company_name, file_type, content]):
            raise HTTPException(status_code=400, detail="Eksik parametre")
        
        # Firma klasörünü oluştur
        company_dir = COMPANIES_DIR / company_name
        company_dir.mkdir(exist_ok=True)
        
        # Excel ise iş emri numarasını çıkar
        work_order_no = None
        if file_type == "excel" and "sheets" in content:
            work_order_no = extract_work_order_number(content["sheets"])
        
        # Hedef klasör yolu
        if work_order_no:
            # İş emri no varsa alt klasör oluştur
            target_dir = company_dir / work_order_no
            target_dir.mkdir(parents=True, exist_ok=True)
            print(f"Klasör oluşturuldu: {target_dir}")
        else:
            # İş emri no yoksa doğrudan firma klasörüne kaydet
            target_dir = company_dir
            print(f"Doğrudan firma klasörüne kaydediliyor")
        
        # Dosya adına firma ismini ekle (büyük harflerle)
        company_prefix = company_name.upper().replace(' ', '_')
        new_filename = f"{company_prefix}_{filename}"
        
        # Hedef dosya yolu
        target_path = target_dir / new_filename
        
        # Dosya tipine göre kaydet
        if file_type == "excel":
            save_excel(target_path, content, filename)
        elif file_type == "word":
            save_word(target_path, content)
        else:
            raise HTTPException(status_code=400, detail="Desteklenmeyen dosya tipi")
        
        return {
            "message": "Dosya başarıyla kaydedildi",
            "path": str(target_path),
            "company": company_name,
            "work_order_no": work_order_no,
            "filename": new_filename
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya kaydetme hatası: {str(e)}")

def save_excel(file_path: Path, content: Dict[str, Any], template_filename: str):
    """
    Excel dosyasını kaydet.
    Şablon dosyayı birebir kopyalar, sadece hücre değerlerini günceller.
    Format/stil bilgileri şablondan korunur.
    """
    template_path = TEMPLATE_DIR / template_filename

    # Şablon dosyayı hedefe kopyala (formatı korumak için)
    if template_path.exists():
        shutil.copy(template_path, file_path)
    else:
        # Şablon yoksa yeni dosya oluştur
        if not file_path.parent.exists():
            file_path.parent.mkdir(parents=True, exist_ok=True)
        wb = openpyxl.Workbook()
        wb.save(file_path)

    # Hedef dosyayı aç ve değerleri yaz
    wb = openpyxl.load_workbook(file_path)
    sheets = content.get("sheets", {})

    for sheet_name, sheet_data in sheets.items():
        if sheet_name not in wb.sheetnames:
            # Şablonda olmayan sheet'i atla (Gerekirse burada yeni sheet de oluşturulabilir)
            continue
            
        ws = wb[sheet_name]
        data = sheet_data.get("data", [])
        
        # Hücre verilerini yaz
        for row_idx, row in enumerate(data):
            for col_idx, cell_data in enumerate(row):
                # Formül varsa formülü yaz
                if "formula" in cell_data:
                    ws.cell(row=row_idx + 1, column=col_idx + 1).value = cell_data["formula"]
                else:
                    # Değeri al
                    value = cell_data.get("value", "")
                    
                    # Sayısal değerleri korumaya çalış (int/float)
                    if isinstance(value, str) and value.strip():
                        # Boşlukları temizleyip sayı kontrolü yap
                        val_str = value.strip()
                        try:
                            if '.' in val_str:
                                value = float(val_str)
                            else:
                                value = int(val_str)
                        except ValueError:
                            # Sayı değilse string olarak bırak
                            pass
                    
                    # Excel hücresine yaz
                    if value != "": # Boş olmayanları yaz
                        ws.cell(row=row_idx + 1, column=col_idx + 1).value = value
        
        # Resimleri kaydet (Place in Cell olanları Place Over Cells'e dönüştürür)
        images = sheet_data.get("images", [])
        if images:
            from openpyxl.drawing.image import Image
            import io
            import base64
            
            for img_data in images:
                try:
                    # Anchor kontrolü
                    anchor = img_data.get("anchor")
                    if not anchor:
                        continue
                        
                    # Base64 verisini çöz
                    b64_data = img_data.get("data")
                    if not b64_data:
                        continue
                        
                    img_bytes = base64.b64decode(b64_data)
                    img_stream = io.BytesIO(img_bytes)
                    
                    # Resmi oluştur
                    img = Image(img_stream)
                    
                    # Resmi hücreye yerleştir
                    img.anchor = anchor
                    
                    # Boyutları ayarla (opsiyonel, orijinal boyut korunabilir)
                    if img_data.get("width"):
                        img.width = img_data["width"]
                    if img_data.get("height"):
                        img.height = img_data["height"]
                        
                    ws.add_image(img)
                    
                except Exception as e:
                    print(f"Resim kaydetme hatası: {str(e)}")
                    continue

    # Aktif sheet'i koru
    active_sheet = content.get("active_sheet")
    if active_sheet and active_sheet in wb.sheetnames:
        wb.active = wb[active_sheet]

    wb.save(file_path)

def save_word(file_path: Path, content: Dict[str, Any]):
    """Word dosyasını kaydet - stil bilgileriyle birlikte"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Paragrafları ekle
    paragraphs = content.get("paragraphs", [])
    for para_data in paragraphs:
        para = doc.add_paragraph()
        
        # Paragraph stilini ayarla
        style_name = para_data.get("style", "Normal")
        try:
            para.style = style_name
        except:
            para.style = "Normal"
        
        # Metni ekle
        text = para_data.get("text", "")
        
        # Format bilgileri varsa uygula
        if "format" in para_data:
            format_data = para_data["format"]
            run = para.add_run(text)
            
            # Font stilleri
            if format_data.get("bold"):
                run.bold = True
            if format_data.get("italic"):
                run.italic = True
            if format_data.get("underline"):
                run.underline = True
            
            # Font boyutu
            if format_data.get("font_size"):
                run.font.size = Pt(format_data["font_size"])
            
            # Font adı
            if format_data.get("font_name"):
                run.font.name = format_data["font_name"]
            
            # Font rengi
            if format_data.get("font_color"):
                try:
                    # RGB hex string'den RGBColor'a çevir
                    rgb_str = format_data["font_color"]
                    if len(rgb_str) == 8:  # ARGB formatı
                        rgb_str = rgb_str[2:]  # Alpha'yı atla
                    if len(rgb_str) == 6:
                        r = int(rgb_str[0:2], 16)
                        g = int(rgb_str[2:4], 16)
                        b = int(rgb_str[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                except:
                    pass
        else:
            # Format bilgisi yoksa sadece metni ekle
            para.add_run(text)
        
        # Hizalama
        if "alignment" in para_data:
            align_str = para_data["alignment"]
            if "LEFT" in align_str:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif "CENTER" in align_str:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "RIGHT" in align_str:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif "JUSTIFY" in align_str:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Tabloları ekle
    tables = content.get("tables", [])
    for table_data in tables:
        if not table_data or not table_data[0]:
            continue
        
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row_data in enumerate(table_data):
            for j, cell_data in enumerate(row_data):
                # Cell_data bir dict veya string olabilir
                if isinstance(cell_data, dict):
                    cell_text = cell_data.get("text", "")
                    cell = table.cell(i, j)
                    cell.text = cell_text
                    
                    # Format bilgileri varsa uygula
                    if "format" in cell_data:
                        format_data = cell_data["format"]
                        if cell.paragraphs:
                            para = cell.paragraphs[0]
                            run = para.runs[0] if para.runs else para.add_run(cell_text)
                            
                            # Font stilleri
                            if format_data.get("bold"):
                                run.bold = True
                            if format_data.get("italic"):
                                run.italic = True
                            
                            # Font boyutu
                            if format_data.get("font_size"):
                                run.font.size = Pt(format_data["font_size"])
                            
                            # Font adı
                            if format_data.get("font_name"):
                                run.font.name = format_data["font_name"]
                            
                            # Font rengi
                            if format_data.get("font_color"):
                                try:
                                    rgb_str = format_data["font_color"]
                                    if len(rgb_str) == 8:
                                        rgb_str = rgb_str[2:]
                                    if len(rgb_str) == 6:
                                        r = int(rgb_str[0:2], 16)
                                        g = int(rgb_str[2:4], 16)
                                        b = int(rgb_str[4:6], 16)
                                        run.font.color.rgb = RGBColor(r, g, b)
                                except:
                                    pass
                else:
                    # Eski format (sadece string)
                    table.cell(i, j).text = str(cell_data)
    
    doc.save(file_path)

def extract_deep_images(file_path: Path) -> Dict[str, List[Dict[str, Any]]]:
    """
    Excel dosyasını ZIP olarak açıp "Place in Cell" (Hücre İçi) resimlerini yakalamaya çalışır.
    Bu yöntem openpyxl'in desteklemediği Rich Value yapısını manuel parse eder.
    """
    import base64
    
    results = {}
    
    try:
        if not zipfile.is_zipfile(file_path):
            return {}

        with zipfile.ZipFile(file_path, 'r') as z:
            # 1. Sheet isimlerini ve yollarını bul (workbook.xml)
            sheet_map = {} # { "rId1": "Sheet1" }
            try:
                wb_root = ET.fromstring(z.read('xl/workbook.xml'))
                
                # Namespace ile uğraşmamak için basit tag kontrolü
                def clean_tag(tag):
                    return tag.split('}')[-1] if '}' in tag else tag

                for sheet in wb_root.findall(".//{*}sheet"):
                    name = sheet.get('name')
                    rId = sheet.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if name and rId:
                        sheet_map[rId] = name
            except:
                pass
            
            # 2. Sheet yollarını bul (_rels/workbook.xml.rels)
            sheet_paths = {} # { "Sheet1": "xl/worksheets/sheet1.xml" }
            try:
                rels_root = ET.fromstring(z.read('xl/_rels/workbook.xml.rels'))
                for rel in rels_root.findall(".//{*}Relationship"):
                    rId = rel.get('Id')
                    target = rel.get('Target')
                    if rId in sheet_map:
                        sheet_name = sheet_map[rId]
                        # Target bazen "worksheets/sheet1.xml", bazen "/xl/worksheets..." olabilir
                        if target.startswith('/'):
                            path = target[1:]
                        elif target.startswith('worksheets'):
                            path = f"xl/{target}"
                        else:
                            path = f"xl/{target}"
                        sheet_paths[sheet_name] = path
            except:
                pass
            
            # 3. Metadata haritasını çıkar (metadata.xml) -> VM index to RV index
            # vm="1" -> richValue index X
            metadata_map = [] 
            try:
                if 'xl/metadata.xml' in z.namelist():
                    meta_root = ET.fromstring(z.read('xl/metadata.xml'))
                    # valueMetadata bloğunu bul
                    # Genellikle: <valueMetadata ...><bk><rc t="1" v="0"/></bk>...</valueMetadata>
                    # rc t="1" demek Rich Value, v="0" ise rdValues.xml'deki index
                    for bk in meta_root.findall(".//{*}valueMetadata/{*}bk"):
                        rc = bk.find(".//{*}rc")
                        if rc is not None and rc.get('t') == '1': # t=1 Rich Value demek
                            metadata_map.append(int(rc.get('v')))
                        else:
                            metadata_map.append(None)
            except:
                print("Metadata okunamadi")

            # 4. Rich Values -> Resim İlişkisi (rdValues.xml ve richValueRel.xml)
            # Basitleştirme: richValueRel.xml genellikle rv indexine karşılık gelen rId'yi tutar
            rv_to_rid = {}
            try:
                if 'xl/richData/richValueRel.xml' in z.namelist():
                    # <richValueRel ...><rel r:id="rId1"/> ... </richValueRel>
                    # Buradaki sıra rdValues.xml'deki rv sırasıyla (veya metadata v indexiyle) aynıdır
                    rv_rel_root = ET.fromstring(z.read('xl/richData/richValueRel.xml'))
                    for idx, rel in enumerate(rv_rel_root.findall(".//{*}rel")):
                        rId = rel.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if rId:
                            rv_to_rid[idx] = rId
            except:
                pass
            
            # 5. rId -> Resim Dosyası Yolu (richData/_rels/richValueRel.xml.rels)
            rid_to_path = {}
            try:
                rels_path = 'xl/richData/_rels/richValueRel.xml.rels'
                if rels_path in z.namelist():
                    rv_rels_root = ET.fromstring(z.read(rels_path))
                    for rel in rv_rels_root.findall(".//{*}Relationship"):
                        rId = rel.get('Id')
                        target = rel.get('Target')
                        if rId and target:
                            # Target genelde "../media/image1.jpg" şeklindedir
                            # Biz xl/richData/ konumundayız, ../ sonrası xl/media olur
                            if target.startswith('../'):
                                full_path = "xl" + target[2:]
                            else:
                                full_path = "xl/richData/" + target
                            rid_to_path[rId] = full_path
            except:
                pass
                
            # 6. Her sheet'i tara ve resimli hücreleri bul
            for sheet_name, xml_path in sheet_paths.items():
                if xml_path not in z.namelist():
                    continue
                
                sheet_images = []
                try:
                    sheet_root = ET.fromstring(z.read(xml_path))
                    # <c r="A1" t="e" vm="1">
                    for cell in sheet_root.findall(".//{*}c"):
                        vm = cell.get('vm')
                        if vm:
                            try:
                                vm_idx = int(vm)
                                # vm 1-based index (genelde), metadata listesi 0-based
                                # Ancak Excel vm indexleri bazen 1'den baslar. Deneme yanilma: genelde vm-1
                                if vm_idx > 0 and (vm_idx - 1) < len(metadata_map):
                                    rv_idx = metadata_map[vm_idx - 1]
                                    
                                    if rv_idx is not None and rv_idx in rv_to_rid:
                                        rId = rv_to_rid[rv_idx]
                                        
                                        if rId in rid_to_path:
                                            image_path = rid_to_path[rId]
                                            
                                            # Resmi zip'ten oku
                                            if image_path in z.namelist():
                                                img_data = z.read(image_path)
                                                base64_data = base64.b64encode(img_data).decode('utf-8')
                                                
                                                # Dosya uzantısını bul
                                                ext = os.path.splitext(image_path)[1][1:].lower()
                                                if ext == 'jpeg': ext = 'jpg'
                                                
                                                # Anchor (A1 gibi)
                                                anchor = cell.get('r')
                                                
                                                sheet_images.append({
                                                    "anchor": anchor,
                                                    "data": base64_data,
                                                    "format": ext,
                                                    "width": 100, # Varsayılan
                                                    "height": 100,
                                                    "type": "in_cell" # İşaretleyici
                                                })
                            except:
                                continue
                except Exception as e:
                    print(f"Sheet parse error ({sheet_name}): {e}")
                
                if sheet_images:
                    results[sheet_name] = sheet_images
                    
    except Exception as e:
        print(f"Deep zip parse error: {e}")
        import traceback
        traceback.print_exc()
        
    return results

@router.get("/search")
async def search_files(query: str):
    """
    Firma dosyaları içerisinde arama yap (İş emri no, parça no, parça adı vb.)
    """
    if not query or len(query) < 2:
        return {"results": [], "count": 0}
    
    query = query.lower().strip()
    results = []
    
    print(f"\n[SEARCH] Arama başlatıldı: '{query}'")
    print(f"[SEARCH] Tarama dizini: {COMPANIES_DIR.absolute()}")
    
    try:
        if not COMPANIES_DIR.exists():
            print(f"[SEARCH] HATA: {COMPANIES_DIR} dizini bulunamadı!")
            return {"results": [], "count": 0}

        # Tüm firmaları ve dosyalarını tara
        for company_dir in COMPANIES_DIR.iterdir():
            if not company_dir.is_dir():
                continue
                
            company_name = company_dir.name
            # print(f"[SEARCH] Firma inceleniyor: {company_name}")
            
            # Alt klasörler dahil tüm dosyaları tara
            for file_path in company_dir.rglob("*"):
                if not file_path.is_file():
                    continue
                
                extension = file_path.suffix.lower()
                if extension not in ['.xlsx', '.xls', '.docx', '.doc']:
                    continue
                
                # Dosya adında ara
                match_reason = None
                if query in file_path.name.lower():
                    match_reason = "Dosya adı eşleşti"
                
                # Dosya içeriğinde ara (eğer henüz eşleşmediyse)
                content_match = False
                if not match_reason:
                    if extension in ['.xlsx', '.xls']:
                        content_match = search_in_excel(file_path, query)
                    elif extension in ['.docx', '.doc']:
                        content_match = search_in_word(file_path, query)
                
                if content_match:
                    match_reason = "Dosya içeriği eşleşti"
                
                if match_reason:
                    print(f"[SEARCH] Eşleşme bulundu: {file_path.name} ({match_reason})")
                    # Relative path'i al (firma klasörüne göre)
                    try:
                        rel_path = file_path.relative_to(COMPANIES_DIR / company_name)
                        results.append({
                            "filename": file_path.name,
                            "company": company_name,
                            "path": str(rel_path).replace('\\', '/'),
                            "reason": match_reason,
                            "extension": extension
                        })
                    except Exception as e:
                        print(f"[SEARCH] Path hatası: {e}")
        
        print(f"[SEARCH] Arama tamamlandı. {len(results)} sonuç bulundu.")
        return {"results": results, "count": len(results)}
    except Exception as e:
        print(f"[SEARCH] Genel arama hatası: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Arama sırasında hata oluştu: {str(e)}")

def search_in_excel(file_path: Path, query: str) -> bool:
    """Excel içeriğinde arama yap"""
    try:
        # data_only=True bazen kaydedilmemiş formüllerde sorun çıkarabiliyor.
        # read_only=True ise bazı büyük dosyalarda daha hızlıdır.
        # İkisini de kaldırarak en güvenli (ama biraz daha yavaş) okumayı deneyelim.
        wb = openpyxl.load_workbook(file_path, data_only=False, read_only=False)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            # Tüm hücreleri tara
            for row in ws.iter_rows():
                for cell in row:
                    cell_value = cell.value
                    if cell_value and query in str(cell_value).lower():
                        return True
        return False
    except Exception as e:
        # print(f"[SEARCH] Excel hata ({file_path.name}): {e}")
        return False

def search_in_word(file_path: Path, query: str) -> bool:
    """Word içeriğinde arama yap"""
    try:
        doc = Document(file_path)
        # Paragraflarda ara
        for para in doc.paragraphs:
            if query in para.text.lower():
                return True
        # Tablolarda ara
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if query in cell.text.lower():
                        return True
        return False
    except Exception as e:
        # print(f"[SEARCH] Word hata ({file_path.name}): {e}")
        return False
